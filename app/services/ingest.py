import os
import fitz        # PyMuPDF
import docx
from typing import Tuple
from app.utils.text import normalize_space

def read_pdf(file_bytes: bytes) -> Tuple[str, int]:
    text_parts = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            text_parts.append(page.get_text("text"))
        pages = len(doc)
    return "\n".join(text_parts), pages

def read_docx(file_bytes: bytes) -> Tuple[str, int]:
    tmp = "_tmp_upload.docx"
    with open(tmp, "wb") as f:
        f.write(file_bytes)
    try:
        d = docx.Document(tmp)
        parts = [p.text for p in d.paragraphs]
        for t in d.tables:
            for r in t.rows:
                parts.append(" | ".join(c.text for c in r.cells))
        return "\n".join(parts), 1
    finally:
        try:
            os.remove(tmp)
        except Exception:
            pass

def load_text_and_pages(filename: str, file_bytes: bytes) -> Tuple[str, int]:
    ext = os.path.splitext((filename or "").lower())[1]
    if ext == ".pdf":
        raw_text, pages = read_pdf(file_bytes)
    elif ext == ".docx":
        raw_text, pages = read_docx(file_bytes)
    else:
        raise ValueError("Formato no soportado. Use PDF o DOCX.")
    return normalize_space(raw_text), pages
