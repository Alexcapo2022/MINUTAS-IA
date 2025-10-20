# app.py
import os, re, json, hashlib
from typing import Dict, Any, Tuple, Optional, List

from fastapi import FastAPI, UploadFile, File, HTTPException
from pydantic import BaseModel, Field
from dotenv import load_dotenv

import fitz           # PyMuPDF
import docx
from openai import OpenAI
from openai import APIConnectionError, RateLimitError, APITimeoutError

# ===================== Config =====================
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
OPENAI_TIMEOUT = int(os.getenv("OPENAI_TIMEOUT", "30"))
if not OPENAI_API_KEY:
    raise RuntimeError("Falta OPENAI_API_KEY en .env")

client = OpenAI(api_key=OPENAI_API_KEY)

app = FastAPI(title="MVP Minutas (secuencial + evidencia)", version="1.3.0")

# ===================== Modelos (Pydantic) =====================
class EvidenceItem(BaseModel):
    evidence_text: str = ""
    char_span: List[int] = Field(default_factory=list)  # [start, end] en el texto plano

class Domicilio(BaseModel):
    direccion: str = ""
    ubigeo: str = ""
    distrito: str = ""
    provincia: str = ""
    departamento: str = ""

class DocumentoAdicional(BaseModel):
    tipo: str = ""           # "Pasaporte", "CE", "RUC", etc.
    numero: str = ""
    observaciones: str = ""

class Persona(BaseModel):
    # datos base
    nombres: str = ""
    apellido_paterno: str = ""
    apellido_materno: str = ""
    nacionalidad: str = ""
    tipo_documento: str = ""     # principal (prioridad DNI)
    numero_documento: str = ""
    docs_adicionales: List[DocumentoAdicional] = Field(default_factory=list)
    profesion_ocupacion: str = ""
    estado_civil: str = ""
    domicilio: Domicilio = Domicilio()
    # opcional para casos no clasificados aún
    role: str = ""  # "otorgante" | "beneficiario" | "indeterminado" | ""
    # evidencia por campo
    evidence: Dict[str, EvidenceItem] = Field(default_factory=dict)

class GeneralesLey(BaseModel):
    otorgantes: List[Persona] = Field(default_factory=list)
    beneficiarios: List[Persona] = Field(default_factory=list)
    indeterminados: List[Persona] = Field(default_factory=list)

class ExtractMapped(BaseModel):
    acto: str = "PODER"
    generales_ley: GeneralesLey = GeneralesLey()
    fecha_minuta: str = ""  # YYYY-MM-DD si es posible
    confidence: Dict[str, Any] = Field(default_factory=lambda: {
        "clasificacion_acto": 0.0,
        "campos": {}
    })
    raw_text_hash: str = ""

class ExtractOut(BaseModel):
    ok: bool
    filename: str
    extension: str
    pages: int
    bytes_size: int
    text_preview: str
    text_hash_sha256: str
    is_poder_guess: float
    mapped: ExtractMapped

# ===================== Utilidades I/O =====================
def normalize_space(t: str) -> str:
    return re.sub(r"\s+", " ", t or "").strip()

def read_pdf(file_bytes: bytes) -> Tuple[str, int]:
    text_parts = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            text_parts.append(page.get_text("text"))
        pages = len(doc)
    return "\n".join(text_parts), pages

def read_docx(file_bytes: bytes) -> Tuple[str, int]:
    tmp = "_tmp_upload.docx"
    with open(tmp, "wb") as f:
        f.write(file_bytes)
    try:
        d = docx.Document(tmp)
        parts = [p.text for p in d.paragraphs]
        for t in d.tables:
            for r in t.rows:
                parts.append(" | ".join(c.text for c in r.cells))
        return "\n".join(parts), 1  # docx no expone páginas; 1 simbólico
    finally:
        try:
            os.remove(tmp)
        except Exception:
            pass

def classify_is_poder(text: str) -> float:
    keys = ["acto notarial", "poder", "escritura pública", "escritura publica"]
    hits = sum(k in text.lower() for k in keys)
    return min(1.0, hits / 3)  # 0, 0.33, 0.66, 1.0

# ===================== Fechas ES → ISO =====================
MONTHS = {
    "enero": "01", "febrero": "02", "marzo": "03", "abril": "04",
    "mayo": "05", "junio": "06", "julio": "07", "agosto": "08",
    "septiembre": "09", "setiembre": "09", "octubre": "10", "noviembre": "11", "diciembre": "12"
}
def to_iso_date(text: str) -> Tuple[str, float]:
    if not text:
        return "", 0.0
    m = re.search(r"\b([0-3]?\d)[/.\-]([01]?\d)[/.\-]((?:19|20)\d{2})\b", text)
    if m:
        d, mo, y = m.groups()
        return f"{y}-{str(mo).zfill(2)}-{str(d).zfill(2)}", 0.9
    m = re.search(r"\b([0-3]?\d)\s+de\s+([a-záéíóú]+)\s+de\s+((?:19|20)\d{2})\b", text, flags=re.IGNORECASE)
    if m:
        d, mon, y = m.groups()
        mon_num = MONTHS.get(mon.lower())
        if mon_num:
            return f"{y}-{mon_num}-{str(int(d)).zfill(2)}", 0.85
    m = re.search(r"\b([a-záéíóú]+)\s+([0-3]?\d),\s*((?:19|20)\d{2})\b", text, flags=re.IGNORECASE)
    if m:
        mon, d, y = m.groups()
        mon_num = MONTHS.get(mon.lower())
        if mon_num:
            return f"{y}-{mon_num}-{str(int(d)).zfill(2)}", 0.75
    return "", 0.0

# ===================== Sanitizadores =====================
def _ensure_evidence_item(v: Any) -> EvidenceItem:
    if isinstance(v, dict):
        et = v.get("evidence_text", "")
        cs = v.get("char_span", [])
        if not isinstance(et, str):
            et = ""
        if isinstance(cs, list) and len(cs) == 2 and all(isinstance(x, int) for x in cs):
            pass
        else:
            cs = []
        return EvidenceItem(evidence_text=et, char_span=cs)
    return EvidenceItem()

def _ensure_evidence_map(value: Any) -> Dict[str, EvidenceItem]:
    out: Dict[str, EvidenceItem] = {}
    if isinstance(value, dict):
        for k, v in value.items():
            out[k] = _ensure_evidence_item(v)
    return out

def _ensure_docs_adicionales(value: Any) -> List[Dict[str, str]]:
    def norm_item(it: Any) -> Dict[str, str]:
        base = {"tipo": "", "numero": "", "observaciones": ""}
        if not isinstance(it, dict):
            return base
        for k in base.keys():
            v = it.get(k, "")
            base[k] = v if isinstance(v, str) else ""
        return base

    if isinstance(value, list):
        return [norm_item(v) for v in value]
    if isinstance(value, dict):
        return [norm_item(value)]
    return []

def _ensure_person(obj: Any) -> Dict[str, Any]:
    base = {
        "nombres": "", "apellido_paterno": "", "apellido_materno": "",
        "nacionalidad": "", "tipo_documento": "", "numero_documento": "",
        "docs_adicionales": [],
        "profesion_ocupacion": "", "estado_civil": "",
        "domicilio": {"direccion": "", "ubigeo": "", "distrito": "", "provincia": "", "departamento": ""},
        "role": "",
        "evidence": {}
    }
    if not isinstance(obj, dict):
        return base
    dom = obj.get("domicilio", {})
    if not isinstance(dom, dict):
        dom = {}
    for k in ["nombres","apellido_paterno","apellido_materno","nacionalidad",
              "tipo_documento","numero_documento","profesion_ocupacion","estado_civil","role"]:
        v = obj.get(k, "")
        base[k] = v if isinstance(v, str) else base[k]
    base["docs_adicionales"] = _ensure_docs_adicionales(obj.get("docs_adicionales", []))
    for k in ["direccion","ubigeo","distrito","provincia","departamento"]:
        v = dom.get(k, "")
        base["domicilio"][k] = v if isinstance(v, str) else base["domicilio"][k]
    base["evidence"] = _ensure_evidence_map(obj.get("evidence", {}))
    return base

def _ensure_person_list(value: Any) -> List[Dict[str, Any]]:
    if isinstance(value, list):
        return [_ensure_person(v) for v in value if isinstance(v, (dict,))]
    if isinstance(value, dict):
        return [_ensure_person(value)]
    return []

def sanitize_llm_out(d: Any, fallback_text_hash: str) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    if not isinstance(d, dict):
        d = {}

    out["acto"] = d.get("acto", "PODER")

    gl = d.get("generales_ley", {})
    if not isinstance(gl, dict):
        gl = {}
    out["generales_ley"] = {
        "otorgantes": _ensure_person_list(gl.get("otorgantes", gl.get("otorgante", []))),
        "beneficiarios": _ensure_person_list(gl.get("beneficiarios", gl.get("beneficiario", []))),
        "indeterminados": _ensure_person_list(gl.get("indeterminados", [])),
    }

    # fecha_minuta top-level (normalizar)
    fecha = d.get("fecha_minuta", "")
    if not isinstance(fecha, str):
        if isinstance(gl.get("fecha_minuta", ""), str):
            fecha = gl.get("fecha_minuta")
        else:
            fecha = ""
    iso, _c = to_iso_date(fecha)
    out["fecha_minuta"] = iso or (fecha if isinstance(fecha, str) else "")

    conf = d.get("confidence", {})
    if not isinstance(conf, dict):
        conf = {}
    campos = conf.get("campos", {})
    if not isinstance(campos, dict):
        campos = {}
    conf["campos"] = campos
    conf.setdefault("clasificacion_acto", 0.0)
    out["confidence"] = conf

    out["raw_text_hash"] = d.get("raw_text_hash", fallback_text_hash)

    # --- Prioridad de documento principal ---
    for role in ["otorgantes", "beneficiarios", "indeterminados"]:
        plist = out["generales_ley"].get(role, [])
        for p in plist:
            main_td = p.get("tipo_documento", "")
            main_nd = p.get("numero_documento", "")
            add = p.get("docs_adicionales", [])

            if (not main_td or not main_nd) and isinstance(add, list):
                dni_idx = next((i for i,x in enumerate(add) if isinstance(x, dict) and x.get("tipo","").upper().startswith("DNI")), None)
                if dni_idx is not None:
                    cand = add.pop(dni_idx)
                    p["tipo_documento"] = cand.get("tipo","")
                    p["numero_documento"] = cand.get("numero","")
                elif add:
                    cand = add.pop(0)
                    p["tipo_documento"] = p.get("tipo_documento","") or cand.get("tipo","")
                    p["numero_documento"] = p.get("numero_documento","") or cand.get("numero","")
            if isinstance(add, list) and main_td.lower().startswith("pasaporte"):
                dni_idx = next((i for i,x in enumerate(add) if isinstance(x, dict) and x.get("tipo","").upper().startswith("DNI")), None)
                if dni_idx is not None:
                    cand = add.pop(dni_idx)
                    add.append({"tipo": main_td, "numero": main_nd, "observaciones": ""})
                    p["tipo_documento"] = cand.get("tipo","")
                    p["numero_documento"] = cand.get("numero","")
            p["docs_adicionales"] = _ensure_docs_adicionales(add)

    return out

# ===================== Post-validaciones =====================
def post_validate(payload: Dict[str, Any]) -> Dict[str, Any]:
    campos = payload.setdefault("confidence", {}).setdefault("campos", {})

    # Valida DNIs/RUCs de TODOS los grupos
    for role in ["otorgantes", "beneficiarios", "indeterminados"]:
        for idx, p in enumerate(payload.get("generales_ley", {}).get(role, [])):
            ndoc = p.get("numero_documento", "")
            key = f"generales_ley.{role}[{idx}].numero_documento"
            if ndoc:
                if re.fullmatch(r"\d{8}", ndoc) or re.fullmatch(r"\d{11}", ndoc):
                    campos[key] = max(campos.get(key, 0.7), 0.9)
                else:
                    campos[key] = min(campos.get(key, 0.6), 0.4)
            # docs adicionales
            for jdx, ad in enumerate(p.get("docs_adicionales", [])):
                num = ad.get("numero","")
                key2 = f"generales_ley.{role}[{idx}].docs_adicionales[{jdx}].numero"
                if num:
                    ok_num = (
                        (ad.get("tipo","").upper().startswith("DNI") and re.fullmatch(r"\d{8}", num)) or
                        (ad.get("tipo","").lower().startswith("ruc") and re.fullmatch(r"\d{11}", num)) or
                        (ad.get("tipo","").lower().startswith("pasaporte"))
                    )
                    campos[key2] = max(campos.get(key2, 0.6), 0.85) if ok_num else min(campos.get(key2, 0.6), 0.4)

    # Fecha ISO
    fecha = payload.get("fecha_minuta", "")
    if fecha and re.fullmatch(r"\d{4}-\d{2}-\d{2}", fecha):
        campos["fecha_minuta"] = max(campos.get("fecha_minuta", 0.7), 0.9)

    return payload

# ===================== Prompt (secuencial + few-shots) =====================
SYSTEM_PROMPT = (
    "Eres un extractor notarial peruano. Lee el texto en orden y arma personas de forma secuencial: "
    "cuando aparece un nombre completo abres una persona activa; mientras sigan atributos contiguos "
    "(documento, nacionalidad, domicilio, estado civil, profesión), los agregas a esa misma persona. "
    "Cierra la persona cuando aparezca un nuevo nombre o la bisagra de rol. "
    "NO inventes datos (si falta, usa \"\"). Devuelve EXCLUSIVAMENTE un único objeto JSON que sigue "
    "exactamente el json_schema. Incluye por cada campo poblado un registro de evidencia "
    "en persona.evidence['ruta.campo'] = {evidence_text, char_span} con índices sobre el texto plano."
    "\nReglas clave:\n"
    "1) Roles: si aparece 'otorga … a favor de …', todo lo anterior es otorgante(s) y lo posterior beneficiario(s). "
    "   Si hay 'y'/comas, separa en múltiples personas. Si no hay bisagra/etiqueta, usa role='indeterminado' y baja confianza.\n"
    "2) Nombres/apellidos (ES-PE): conserva tildes y apóstrofes. En ≥4 tokens, toma los 2 últimos como apellidos; resto = nombres. "
    "   Si solo hay un apellido, úsalo como paterno y deja materno vacío.\n"
    "3) Documentos: prioriza DNI como principal tipo/numero. Otros (Pasaporte/CE/RUC) a docs_adicionales[]. "
    "   Reconoce variantes: 'DNI', 'D.N.I.', 'Documento Nacional de Identidad', 'DNI N.º', 'DNI No.', 'DNI #', 'DNI-Nro.'; "
    "   acepta separadores '/' y '-'.\n"
    "4) Domicilio: las frases 'con domicilio en|domiciliado en|domicilio:' asignan dirección a la persona ACTIVA. "
    "   Si dice 'ambos/todos domiciliados', replica a todas las personas del grupo. Normaliza comas y espacios ("
    "'756 , La Molina' → '756, La Molina').\n"
    "5) fecha_minuta: top-level; convierte a ISO si puedes.\n"
    "6) confidence.campos: 0–1. Sube confianza si DNI=8 dígitos, RUC=11, fecha ISO válida o si aplicaste la bisagra de roles.\n"
)

def build_json_contract() -> Dict[str, Any]:
    base = ExtractMapped().model_dump()
    base.pop("raw_text_hash", None)
    # Arrays explícitos vacíos
    base["generales_ley"]["otorgantes"] = []
    base["generales_ley"]["beneficiarios"] = []
    base["generales_ley"]["indeterminados"] = []
    return base

FEW_SHOTS = """
FEW-SHOTS:

Caso A (otorga … a favor de …, DNI + Pasaporte):
Texto:
"… otorga el señor OLIVER THOMAS ALEXANDER STARK PREUSS, de nacionalidad peruano-alemana, identificado con DNI # 10322575 y pasaporte alemán # C4FNKZ6ZF, con domicilio en Jr. El Golf 756 , La Molina, a favor de WILFREDO MIGUEL YAÑEZ LAZO, identificado con DNI N° 07907927 …"
Pistas:
- Personas en secuencia, domicilio asignado a la persona activa.
- DNI como documento principal, pasaporte en docs_adicionales.
- Coma rara en domicilio se normaliza.
Salida (parcial, forma esperada):
{
  "acto": "PODER",
  "generales_ley": {
    "otorgantes": [
      {
        "nombres": "OLIVER THOMAS ALEXANDER",
        "apellido_paterno": "STARK",
        "apellido_materno": "PREUSS",
        "nacionalidad": "peruano-alemana",
        "tipo_documento": "DNI",
        "numero_documento": "10322575",
        "docs_adicionales": [{"tipo": "Pasaporte", "numero": "C4FNKZ6ZF", "observaciones": "alemán"}],
        "domicilio": {"direccion": "Jr. El Golf 756, La Molina", "ubigeo":"", "distrito":"La Molina", "provincia":"", "departamento":""},
        "role": "otorgante",
        "evidence": {
          "numero_documento": {"evidence_text":"DNI # 10322575","char_span":[100,115]},
          "domicilio.direccion": {"evidence_text":"Jr. El Golf 756 , La Molina","char_span":[120,148]}
        }
      }
    ],
    "beneficiarios": [
      {
        "nombres": "WILFREDO MIGUEL",
        "apellido_paterno": "YAÑEZ",
        "apellido_materno": "LAZO",
        "tipo_documento": "DNI",
        "numero_documento": "07907927",
        "evidence": {
          "numero_documento": {"evidence_text":"DNI N° 07907927","char_span":[180,196]}
        }
      }
    ],
    "indeterminados": []
  },
  "fecha_minuta": "",
  "confidence": {"clasificacion_acto": 0.9, "campos": {}}
}

Caso B (domicilio con Jr. y coma rara, sin bisagra clara):
Texto:
"… MARÍA LOURDES D'BROT VARGAS, con domicilio en Jr. Los Sauces 123 , Surco, identificada con Documento Nacional de Identidad 12345678 …"
Salida (parcial):
{
  "generales_ley": {
    "otorgantes": [
      {
        "nombres": "MARÍA LOURDES",
        "apellido_paterno": "D'BROT",
        "apellido_materno": "VARGAS",
        "tipo_documento": "DNI",
        "numero_documento": "12345678",
        "domicilio": {"direccion":"Jr. Los Sauces 123, Surco", "distrito":"Santiago de Surco"},
        "role":"indeterminado"
      }
    ],
    "beneficiarios": [],
    "indeterminados": []
  }
}
"""

def call_llm_extract(text: str) -> Dict[str, Any]:
    json_contract = build_json_contract()
    user_prompt = (
        "json_schema:\n"
        f"{json.dumps(json_contract, ensure_ascii=False)}\n\n"
        f"{FEW_SHOTS}\n\n"
        "CONTENIDO (texto plano normalizado):\n"
        f"{text}"
    )
    resp = client.chat.completions.create(
        model=OPENAI_MODEL,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0,
        timeout=OPENAI_TIMEOUT,
        response_format={"type": "json_object"}
    )
    content = resp.choices[0].message.content
    return json.loads(content)

# ===================== Endpoints =====================
@app.get("/health")
def health():
    return {"status": "ok", "model": OPENAI_MODEL}

@app.get("/")
def root():
    return {"message": "API viva. Sube un PDF/DOCX a /extract y obtén JSON secuencial con evidencia."}

@app.post("/extract", response_model=ExtractOut)
async def extract(file: UploadFile = File(...)):
    raw = await file.read()
    if not raw:
        raise HTTPException(status_code=400, detail="Archivo vacío")

    name = (file.filename or "").lower()
    ext = os.path.splitext(name)[1]

    if ext == ".pdf":
        raw_text, pages = read_pdf(raw)
    elif ext == ".docx":
        raw_text, pages = read_docx(raw)
    else:
        raise HTTPException(status_code=415, detail="Formato no soportado. Use PDF o DOCX.")

    raw_text_norm = normalize_space(raw_text)
    text_hash = hashlib.sha256(raw_text_norm.encode("utf-8")).hexdigest()
    guess = classify_is_poder(raw_text_norm)

    # LLM + sanitizado
    llm_out: Optional[Dict[str, Any]] = None
    last_err = None
    for _ in range(2):
        try:
            llm_out = call_llm_extract(raw_text_norm)
            break
        except (APIConnectionError, RateLimitError, APITimeoutError) as e:
            last_err = e
        except Exception as e:
            last_err = e
            break
    if llm_out is None:
        raise HTTPException(status_code=502, detail=f"Falla LLM: {last_err}")

    llm_out["raw_text_hash"] = text_hash
    llm_out = sanitize_llm_out(llm_out, text_hash)
    llm_out = post_validate(llm_out)

    # Ajuste de confianza por clasificación local
    conf = llm_out.setdefault("confidence", {})
    conf["clasificacion_acto"] = max(conf.get("clasificacion_acto", 0.0), guess)

    # Validación Pydantic final
    try:
        mapped = ExtractMapped(**llm_out)
    except Exception:
        safe = sanitize_llm_out({}, text_hash)
        mapped = ExtractMapped(**safe)

    return ExtractOut(
        ok=True,
        filename=file.filename,
        extension=ext,
        pages=pages,
        bytes_size=len(raw),
        text_preview=raw_text_norm[:600],
        text_hash_sha256=text_hash,
        is_poder_guess=guess,
        mapped=mapped
    )
