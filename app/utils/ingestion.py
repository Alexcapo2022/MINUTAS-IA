# app/utils/ingestion.py
from fastapi import UploadFile, HTTPException
from io import BytesIO

ALLOWED_MIME = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
}

def _ensure_allowed(file: UploadFile):
    if file.content_type not in ALLOWED_MIME:
        raise HTTPException(
            status_code=400,
            detail=f"Formato no permitido: {file.content_type}. Solo PDF o Word (DOC/DOCX)."
        )

async def get_text_from_upload(file: UploadFile) -> str:
    """
    Lee un UploadFile (PDF o Word) y devuelve el texto extraído.
    """
    _ensure_allowed(file)
    raw = await file.read()

    if file.content_type == "application/pdf":
        try:
            from pypdf import PdfReader
        except Exception as e:
            raise HTTPException(500, f"Dependencia faltante pypdf: {e}")
        reader = PdfReader(BytesIO(raw))
        parts = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        text = "\n".join(parts).strip()
        if not text:
            raise HTTPException(400, "No se pudo extraer texto del PDF.")
        return text

    # DOCX/DOC
    try:
        from docx import Document
    except Exception as e:
        raise HTTPException(500, f"Dependencia faltante python-docx: {e}")

    try:
        doc = Document(BytesIO(raw))
    except Exception as e:
        raise HTTPException(400, f"No se pudo abrir el archivo Word: {e}")

    # Extraer párrafos y tablas en orden de aparición en XML
    # Para simplificar y dado que el orden exacto no afecta severamente al LLM,
    # recopilamos párrafos y luego tablas (o mejor, iteramos iter_block_items si tuviéramos la función, 
    # pero a falta de ella extraemos ambos).
    
    text_blocks = []
    
    # Primero insertamos los párrafos normales
    for p in doc.paragraphs:
        if p.text.strip():
            text_blocks.append(p.text.strip())
            
    # Luego recorremos todas las tablas y sacamos el contenido fila por fila
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cleaned_cell = cell.text.strip().replace('\n', ' ')
                if cleaned_cell:
                    row_text.append(cleaned_cell)
            if row_text:
                text_blocks.append(" | ".join(row_text))

    text = "\n".join(text_blocks).strip()
    if not text:
        raise HTTPException(400, "No se encontró texto en el archivo Word.")
    return text
